"""
Admin API endpoints for the management console.
Provides JSON data for Dashboard, Customers, Conversations, Follow-ups, and Bridge Status.
"""
import json
import logging
import os
from datetime import datetime
from pathlib import Path

import httpx
import csv
import io

import openpyxl

logger = logging.getLogger(__name__)
from fastapi import APIRouter, File, Query, UploadFile

from config import get_config, get_data_dir
from database import get_connection, update_customer
from qr_state import get_pending_qr, clear_pending_qr
from bot_state import is_paused, toggle_paused, set_paused

_ocr_instance = None

def _get_ocr():
    global _ocr_instance
    if _ocr_instance is None:
        from rapidocr import RapidOCR
        _ocr_instance = RapidOCR()
    return _ocr_instance

router = APIRouter(prefix="/admin/api")


# ── Setup ──────────────────────────────────────────────────────────────────────

@router.get("/setup-status")
async def setup_status():
    """Check if the app is configured (API key set in .env file)."""
    data_dir = get_data_dir()
    env_path = data_dir / ".env"
    has_key = False
    if env_path.exists():
        for line in env_path.read_text().split("\n"):
            line = line.strip()
            if line.startswith("LLM_API_KEY="):
                val = line.partition("=")[2].strip().strip('"').strip("'")
                if val:
                    has_key = True
                break
    return {"configured": has_key}


@router.post("/setup")
async def save_setup(data: dict):
    """Save initial configuration (API key)."""
    api_key = data.get("api_key", "").strip()
    if not api_key:
        return {"error": "api_key required"}

    data_dir = get_data_dir()
    env_path = data_dir / ".env"

    # Read existing .env or create new one
    existing = {}
    if env_path.exists():
        for line in env_path.read_text().split("\n"):
            line = line.strip()
            if "=" in line and not line.startswith("#"):
                k, _, v = line.partition("=")
                existing[k.strip()] = v.strip().strip('"').strip("'")

    existing["LLM_API_KEY"] = api_key

    # Write back
    lines = [f"{k}={v}" for k, v in existing.items()]
    env_path.write_text("\n".join(lines) + "\n")

    # Set in current process
    os.environ["LLM_API_KEY"] = api_key

    return {"ok": True}


# ── Dashboard ──────────────────────────────────────────────────────────────────

@router.get("/dashboard")
async def dashboard():
    conn = get_connection()
    today = datetime.now().strftime("%Y-%m-%d")
    now_iso = datetime.now().isoformat()

    total = conn.execute("SELECT COUNT(*) FROM customers").fetchone()[0]
    active = conn.execute("SELECT COUNT(*) FROM customers WHERE status='active'").fetchone()[0]

    # Today's stats
    today_ai_replies = conn.execute(
        "SELECT COUNT(*) FROM conversations WHERE ai_generated=1 AND direction='outbound' AND date(sent_at)=?",
        (today,)
    ).fetchone()[0]
    today_inbound = conn.execute(
        "SELECT COUNT(*) FROM conversations WHERE direction='inbound' AND date(sent_at)=?",
        (today,)
    ).fetchone()[0]

    # Follow-up stats
    due_today = conn.execute(
        "SELECT COUNT(*) FROM follow_up_schedule fs JOIN customers c ON fs.customer_id=c.id "
        "WHERE fs.active=1 AND c.status='active' AND date(fs.next_followup_at)=?", (today,)
    ).fetchone()[0]
    due_overdue = conn.execute(
        "SELECT COUNT(*) FROM follow_up_schedule fs JOIN customers c ON fs.customer_id=c.id "
        "WHERE fs.active=1 AND c.status='active' AND fs.next_followup_at <= ?", (now_iso,)
    ).fetchone()[0]
    sent_fu_total = conn.execute("SELECT COUNT(*) FROM sent_followups WHERE status='sent'").fetchone()[0]
    sent_fu_today = conn.execute(
        "SELECT COUNT(*) FROM sent_followups WHERE status='sent' AND date(sent_at)=?", (today,)
    ).fetchone()[0]
    failed_fu_today = conn.execute(
        "SELECT COUNT(*) FROM sent_followups WHERE status='failed' AND date(sent_at)=?", (today,)
    ).fetchone()[0]

    total_schedules = conn.execute("SELECT COUNT(*) FROM follow_up_schedule WHERE active=1").fetchone()[0]
    total_conv = conn.execute("SELECT COUNT(*) FROM conversations").fetchone()[0]

    # Recently failed sends
    failed_recent = conn.execute(
        "SELECT c.name, c.phone, sf.error_message, sf.sent_at FROM sent_followups sf "
        "JOIN customers c ON sf.customer_id=c.id "
        "WHERE sf.status='failed' ORDER BY sf.sent_at DESC LIMIT 5"
    ).fetchall()

    # Recently active customers
    recent_active = conn.execute(
        "SELECT c.id, c.name, c.phone, c.company, c.status, "
        "MAX(conv.sent_at) as last_contact FROM customers c "
        "JOIN conversations conv ON c.id=conv.customer_id "
        "GROUP BY c.id ORDER BY last_contact DESC LIMIT 10"
    ).fetchall()

    # Bridge status
    bridge = await _get_bridge_status()

    # Upcoming schedule
    upcoming = conn.execute(
        "SELECT date(fs.next_followup_at) as dt, COUNT(*) as cnt FROM follow_up_schedule fs "
        "JOIN customers c ON fs.customer_id=c.id "
        "WHERE fs.active=1 AND c.status='active' AND date(fs.next_followup_at) >= ? "
        "GROUP BY dt ORDER BY dt LIMIT 14", (today,)
    ).fetchall()

    return {
        "today": today,
        "counts": {
            "total_customers": total, "active_customers": active,
            "today_ai_replies": today_ai_replies, "today_inbound": today_inbound,
            "due_today": due_today, "due_overdue": due_overdue,
            "sent_fu_total": sent_fu_total, "sent_fu_today": sent_fu_today,
            "failed_fu_today": failed_fu_today,
            "total_schedules": total_schedules, "total_conversations": total_conv,
        },
        "bridge": bridge,
        "failed_recent": [dict(r) for r in failed_recent],
        "recent_active": [
            {"id": r["id"], "name": r["name"], "phone": r["phone"],
             "company": r["company"], "status": r["status"], "last_contact": r["last_contact"]}
            for r in recent_active
        ],
        "upcoming": [{"date": r["dt"], "count": r["cnt"]} for r in upcoming],
    }


# ── Customers ──────────────────────────────────────────────────────────────────

@router.get("/customers")
async def customers(
    page: int = Query(1, ge=1),
    per_page: int = Query(30, ge=5, le=100),
    search: str = Query(""),
    status: str = Query(""),
    sort: str = Query("updated_at"),
    order: str = Query("desc"),
):
    conn = get_connection()
    allowed_sorts = {"name", "company", "status", "created_at", "updated_at"}
    if sort not in allowed_sorts:
        sort = "updated_at"
    if order not in ("asc", "desc"):
        order = "desc"

    where = ["1=1"]
    params = []
    if search:
        where.append("(c.name LIKE ? OR c.phone LIKE ? OR c.company LIKE ?)")
        s = f"%{search}%"
        params.extend([s, s, s])
    if status:
        where.append("c.status = ?")
        params.append(status)

    where_clause = " AND ".join(where)

    total = conn.execute(f"SELECT COUNT(*) FROM customers c WHERE {where_clause}", params).fetchone()[0]

    params_with_limit = params + [per_page, (page - 1) * per_page]
    rows = conn.execute(
        f"""SELECT c.*,
                   (SELECT COUNT(*) FROM conversations WHERE customer_id=c.id) as conv_count,
                   (SELECT MAX(sent_at) FROM conversations WHERE customer_id=c.id) as last_msg,
                   (SELECT next_followup_at FROM follow_up_schedule WHERE customer_id=c.id AND active=1 LIMIT 1) as next_fu
            FROM customers c WHERE {where_clause}
            ORDER BY c.{sort} {order} LIMIT ? OFFSET ?""",
        params_with_limit
    ).fetchall()

    return {
        "total": total,
        "page": page,
        "per_page": per_page,
        "items": [dict(r) for r in rows],
    }


# ── Invalid Phone Detection & Cleanup (must be BEFORE {customer_id} route) ──

@router.get("/customers/invalid-phones")
async def list_invalid_phones():
    """Detect customers whose phone field contains names instead of phone numbers."""
    conn = get_connection()
    rows = conn.execute(
        """SELECT c.id, c.phone, c.name, c.status, c.company, c.notes,
                  (SELECT COUNT(*) FROM conversations WHERE customer_id=c.id) as conv_count,
                  (SELECT COUNT(*) FROM follow_up_schedule WHERE customer_id=c.id AND active=1) as active_fu_count
           FROM customers c
           WHERE c.phone GLOB '*[A-Za-z]*' AND c.phone NOT GLOB '+*'
           ORDER BY c.status, c.id"""
    ).fetchall()

    items = [dict(r) for r in rows]
    total = len(items)
    with_conv = sum(1 for r in items if r["conv_count"] > 0)
    with_fu = sum(1 for r in items if r["active_fu_count"] > 0)
    safe_to_delete = sum(1 for r in items if r["conv_count"] == 0)

    return {
        "total_invalid": total,
        "with_conversations": with_conv,
        "with_active_followup": with_fu,
        "safe_to_delete": safe_to_delete,
        "items": items,
    }


@router.post("/customers/cleanup-invalid-phones")
async def cleanup_invalid_phones(data: dict):
    """Batch cleanup invalid phone numbers.

    Actions:
      - "detect_only": return list without making changes
      - "delete_safe": delete customers with invalid phones AND no conversations
      - "mark_inactive": mark customers with invalid phones as inactive (if they have conversations)
    """
    action = data.get("action", "detect_only")
    conn = get_connection()

    invalid = conn.execute(
        """SELECT c.id, c.phone, c.name, c.status,
                  (SELECT COUNT(*) FROM conversations WHERE customer_id=c.id) as conv_count
           FROM customers c
           WHERE c.phone GLOB '*[A-Za-z]*' AND c.phone NOT GLOB '+*'"""
    ).fetchall()

    result = {"action": action, "total_invalid": len(invalid)}

    if action == "detect_only":
        result["message"] = f"发现 {len(invalid)} 个无效号码。使用 action=delete_safe 或 action=mark_inactive 执行清理。"
        return result

    if action == "delete_safe":
        to_delete = [r for r in invalid if r["conv_count"] == 0]
        deleted_ids = []
        for r in to_delete:
            cid = r["id"]
            conn.execute("DELETE FROM sent_followups WHERE customer_id=?", (cid,))
            conn.execute("DELETE FROM follow_up_schedule WHERE customer_id=?", (cid,))
            conn.execute("DELETE FROM conversations WHERE customer_id=?", (cid,))
            conn.execute("DELETE FROM customers WHERE id=?", (cid,))
            deleted_ids.append({"id": cid, "phone": r["phone"], "name": r["name"]})
        conn.commit()
        result["deleted_count"] = len(deleted_ids)
        result["deleted"] = deleted_ids
        result["skipped_count"] = len(invalid) - len(deleted_ids)
        result["message"] = f"已删除 {len(deleted_ids)} 个无效号码客户（无对话记录），跳过 {result['skipped_count']} 个（有对话记录需手动处理）"
        return result

    if action == "mark_inactive":
        updated_count = 0
        for r in invalid:
            conn.execute(
                "UPDATE customers SET status='inactive', notes=COALESCE(notes,'') || ' [自动标记：电话号码无效]', updated_at=? WHERE id=?",
                (datetime.now().isoformat(), r["id"]))
            conn.execute("UPDATE follow_up_schedule SET active=0 WHERE customer_id=?", (r["id"],))
            updated_count += 1
        conn.commit()
        result["marked_inactive"] = updated_count
        result["message"] = f"已将 {updated_count} 个无效号码客户标记为 inactive，其跟进计划已全部暂停。"
        return result

    result["error"] = f"未知操作: {action}，支持 detect_only / delete_safe / mark_inactive"
    return result


@router.get("/customers/{customer_id}")
async def customer_detail(customer_id: int):
    conn = get_connection()
    row = conn.execute("SELECT * FROM customers WHERE id=?", (customer_id,)).fetchone()
    if not row:
        return {"error": "not found"}

    convs = conn.execute(
        "SELECT * FROM conversations WHERE customer_id=? ORDER BY sent_at DESC LIMIT 50",
        (customer_id,)
    ).fetchall()

    fu = conn.execute(
        "SELECT * FROM follow_up_schedule WHERE customer_id=? AND active=1 LIMIT 1",
        (customer_id,)
    ).fetchone()

    sent_fus = conn.execute(
        "SELECT * FROM sent_followups WHERE customer_id=? ORDER BY sent_at DESC LIMIT 10",
        (customer_id,)
    ).fetchall()

    return {
        "customer": dict(row),
        "conversations": [dict(c) for c in reversed(convs)],
        "followup_schedule": dict(fu) if fu else None,
        "sent_followups": [dict(s) for s in sent_fus],
    }


@router.post("/customers/{customer_id}/status")
async def update_customer_status(customer_id: int, data: dict):
    allowed = {"status", "notes", "tags", "name", "company", "phone", "email"}
    updates = {k: v for k, v in data.items() if k in allowed}
    if not updates:
        return {"error": "no valid fields"}
    row = get_connection().execute("SELECT phone FROM customers WHERE id=?", (customer_id,)).fetchone()
    if not row:
        return {"error": "not found"}
    new_phone = updates.pop("phone", None)
    try:
        result = update_customer(row["phone"], new_phone=new_phone, **updates)
    except Exception as e:
        err_msg = str(e)
        if "UNIQUE constraint" in err_msg and "customers.phone" in err_msg and new_phone:
            # ── Smart merge: reassign data from the phone owner to current customer ──
            existing = get_connection().execute(
                "SELECT * FROM customers WHERE phone=?", (new_phone,)
            ).fetchone()
            if existing and existing["id"] != customer_id:
                conn = get_connection()
                # Reassign all related data from existing customer to current customer
                conn.execute("UPDATE conversations SET customer_id=? WHERE customer_id=?",
                             (customer_id, existing["id"]))
                conn.execute("UPDATE follow_up_schedule SET customer_id=? WHERE customer_id=?",
                             (customer_id, existing["id"]))
                conn.execute("UPDATE sent_followups SET customer_id=? WHERE customer_id=?",
                             (customer_id, existing["id"]))
                conn.execute("UPDATE orders SET customer_id=? WHERE customer_id=?",
                             (customer_id, existing["id"]))
                # Delete the old customer (frees the phone number)
                conn.execute("DELETE FROM customers WHERE id=?", (existing["id"],))
                conn.commit()
                logger.info(
                    f"Merged customer #{existing['id']} ({existing['name']}, {existing['phone']}) "
                    f"into #{customer_id} — all related records reassigned, old record deleted"
                )
                # Now the phone is free — retry the update
                try:
                    result = update_customer(row["phone"], new_phone=new_phone, **updates)
                except Exception as e2:
                    return {"ok": False, "error": f"合并后更新失败：{e2}"}
                return {
                    "ok": True,
                    "customer": result,
                    "merged_from": {
                        "id": existing["id"],
                        "name": existing["name"],
                        "phone": existing["phone"],
                    },
                }
            return {"ok": False, "error": f"电话号码 {new_phone} 已存在，无法保存。请检查是否与其他客户重复。"}
        return {"ok": False, "error": f"保存失败：{err_msg}"}
    if result is None and new_phone:
        return {"ok": False, "error": f"更新失败，电话号码可能已被占用或客户不存在。"}
    return {"ok": True, "customer": result}


NAME_KEYS = ["姓名", "name", "名称"]
PHONE_KEYS = ["电话", "phone", "手机", "手机号", "电话号码", "tel", "mobile"]
EMAIL_KEYS = ["邮箱", "email", "e-mail", "邮件"]
NOTES_KEYS = ["备注", "notes", "note", "说明", "remark"]
INFO_KEYS = ["客户具体信息", "联系人信息", "contact info", "contact"]
COMPANY_KEYS = ["公司名", "公司名、网址", "公司", "company", "公司名称"]


def _parse_info_field(raw: str) -> dict:
    """Parse N:/E:/M: fields from a combined info column like '客户具体信息'."""
    result = {"name": "", "email": "", "phone": ""}
    if not raw:
        return result
    for line in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = line.strip()
        if not line:
            continue
        if ":" in line:
            prefix, _, value = line.partition(":")
            prefix = prefix.strip().upper()
            value = value.strip()
            if prefix == "N" and not result["name"]:
                result["name"] = value
            elif prefix == "E" and not result["email"]:
                result["email"] = value
            elif prefix == "M" and not result["phone"]:
                result["phone"] = value
    return result


def _find_column(headers: list, keys: list) -> str | None:
    """Find the first header that matches one of the given keys (case-insensitive)."""
    key_set = {k.lower() for k in keys}
    for h in headers:
        hl = h.strip().lower()
        if hl in key_set:
            return h
    return None


def _build_col_map(headers: list):
    """Build column mapping from headers. Returns (col_map, has_info_col)."""
    col_map = {}
    has_info_col = False
    for h in headers:
        hl = h.strip().lower()
        if _find_column([h], NAME_KEYS):
            col_map["name"] = h
        elif _find_column([h], PHONE_KEYS):
            col_map["phone"] = h
        elif _find_column([h], EMAIL_KEYS):
            col_map["email"] = h
        elif _find_column([h], NOTES_KEYS):
            col_map["notes"] = h
        elif _find_column([h], COMPANY_KEYS):
            col_map["company"] = h
        elif _find_column([h], INFO_KEYS):
            col_map["info"] = h
            has_info_col = True
    return col_map, has_info_col


@router.post("/customers/import")
async def import_customers(file: UploadFile = File(...)):
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    row_offset = 2  # default for CSV (header row 1, data from row 2)

    if ext == "csv":
        content = (await file.read()).decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(content))
        rows = list(reader)
        if not rows:
            return {"error": "文件为空或无法解析"}
        headers = list(rows[0].keys())
    elif ext in ("xlsx", "xls"):
        content = await file.read()
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        ws = wb.active
        # Read first 5 rows to find header row
        max_scan = min(5, ws.max_row or 5)
        preview_rows = []
        for row in ws.iter_rows(min_row=1, max_row=max_scan, values_only=True):
            preview_rows.append([str(c or "").strip() for c in row])

        # Find best header row: the one with most known column matches
        best_row_idx = 0
        best_score = 0
        all_known = NAME_KEYS + PHONE_KEYS + EMAIL_KEYS + NOTES_KEYS + COMPANY_KEYS + INFO_KEYS
        all_lower = {k.lower() for k in all_known}
        for idx, row_data in enumerate(preview_rows):
            score = sum(1 for c in row_data if c.lower() in all_lower)
            if score > best_score:
                best_score = score
                best_row_idx = idx

        if best_score == 0:
            wb.close()
            return {"error": "未找到可识别的表头列，请确保表格包含姓名/电话/客户具体信息等列"}

        headers = preview_rows[best_row_idx]
        row_offset = best_row_idx + 2  # actual Excel row number of first data row
        rows = [dict(zip(headers, [str(c or "") if c is not None else "" for c in row]))
                for row in ws.iter_rows(min_row=row_offset, values_only=True)]
        wb.close()
    else:
        return {"error": "仅支持 .xlsx 或 .csv 文件"}

    if not rows:
        return {"error": "文件为空或无法解析"}

    col_map, has_info_col = _build_col_map(headers)

    if "phone" not in col_map and not has_info_col:
        return {"error": "未找到电话/手机列，请确保表头包含'电话'或'phone'列，或包含'客户具体信息'列"}

    conn = get_connection()
    imported = 0
    skipped = []

    for i, row in enumerate(rows):
        row_num = i + row_offset
        if has_info_col:
            info_raw = (row.get(col_map.get("info", ""), "") or "").strip()
            parsed = _parse_info_field(info_raw)
            phone_raw = parsed["phone"]
            name = parsed["name"]
            email = parsed["email"]
        else:
            phone_raw = (row.get(col_map.get("phone", ""), "") or "").strip()
            name = (row.get(col_map.get("name", ""), "") or "").strip()
            email = (row.get(col_map.get("email", ""), "") or "").strip()

        # Also try dedicated columns if they exist (override parsed values)
        if "name" in col_map:
            explicit_name = (row.get(col_map["name"], "") or "").strip()
            if explicit_name:
                name = explicit_name
        if "email" in col_map:
            explicit_email = (row.get(col_map["email"], "") or "").strip()
            if explicit_email:
                email = explicit_email

        phone = ''.join(c for c in phone_raw if c.isdigit() or c == '+')
        if not phone:
            skipped.append({"row": row_num, "reason": "电话为空"})
            continue

        if not name:
            skipped.append({"row": row_num, "reason": "姓名为空"})
            continue

        existing = conn.execute("SELECT id FROM customers WHERE phone=?", (phone,)).fetchone()
        if existing:
            skipped.append({"row": row_num, "reason": f"电话 {phone} 已存在"})
            continue

        notes = (row.get(col_map.get("notes", ""), "") or "").strip()
        company = (row.get(col_map.get("company", ""), "") or "").strip()

        try:
            conn.execute(
                "INSERT INTO customers(phone, name, email, company, notes) VALUES(?,?,?,?,?)",
                (phone, name, email, company, notes))
            conn.commit()
            imported += 1
        except Exception as e:
            skipped.append({"row": row_num, "reason": f"写入失败: {str(e)}"})

    return {"imported": imported, "skipped": skipped, "total": len(rows)}


# ── Conversations ──────────────────────────────────────────────────────────────

@router.delete("/conversations/{conversation_id}")
async def delete_conversation(conversation_id: int):
    conn = get_connection()
    row = conn.execute("SELECT id FROM conversations WHERE id=?", (conversation_id,)).fetchone()
    if not row:
        return {"error": "not found"}
    conn.execute("DELETE FROM conversations WHERE id=?", (conversation_id,))
    conn.commit()
    return {"ok": True}


@router.delete("/customers/{customer_id}")
async def delete_customer(customer_id: int):
    """Delete a customer and all related records (conversations, followup schedule, sent followups)."""
    conn = get_connection()
    row = conn.execute("SELECT id, name, phone FROM customers WHERE id=?", (customer_id,)).fetchone()
    if not row:
        return {"ok": False, "error": "客户不存在"}
    conn.execute("DELETE FROM sent_followups WHERE customer_id=?", (customer_id,))
    conn.execute("DELETE FROM follow_up_schedule WHERE customer_id=?", (customer_id,))
    conn.execute("DELETE FROM conversations WHERE customer_id=?", (customer_id,))
    conn.execute("DELETE FROM customers WHERE id=?", (customer_id,))
    conn.commit()
    logger.info(f"Deleted customer #{customer_id}: {row['name']} ({row['phone']}) — all related records removed")
    return {"ok": True, "deleted": {"id": customer_id, "name": row["name"], "phone": row["phone"]}}


@router.delete("/customers/{customer_id}/conversations")
async def delete_customer_conversations(customer_id: int):
    conn = get_connection()
    conn.execute("DELETE FROM conversations WHERE customer_id=?", (customer_id,))
    conn.commit()
    return {"ok": True}


@router.get("/conversations")
async def conversations_list(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=5, le=50),
    customer_id: int = Query(None),
):
    conn = get_connection()
    if customer_id:
        total = conn.execute(
            "SELECT COUNT(*) FROM conversations WHERE customer_id=?", (customer_id,)
        ).fetchone()[0]
        rows = conn.execute(
            "SELECT conv.*, c.name, c.phone FROM conversations conv "
            "JOIN customers c ON conv.customer_id=c.id "
            "WHERE conv.customer_id=? ORDER BY conv.sent_at DESC LIMIT ? OFFSET ?",
            (customer_id, per_page, (page - 1) * per_page)
        ).fetchall()
    else:
        total = conn.execute("SELECT COUNT(*) FROM conversations").fetchone()[0]
        rows = conn.execute(
            "SELECT conv.*, c.name, c.phone FROM conversations conv "
            "JOIN customers c ON conv.customer_id=c.id "
            "ORDER BY conv.sent_at DESC LIMIT ? OFFSET ?",
            (per_page, (page - 1) * per_page)
        ).fetchall()

    return {
        "total": total, "page": page, "per_page": per_page,
        "items": [dict(r) for r in rows],
    }


@router.get("/customers-with-conversations")
async def customers_with_conversations():
    conn = get_connection()
    rows = conn.execute(
        "SELECT c.id, c.name, c.phone, c.company, c.status, "
        "MAX(conv.sent_at) as last_msg, COUNT(conv.id) as msg_count "
        "FROM customers c JOIN conversations conv ON c.id=conv.customer_id "
        "GROUP BY c.id ORDER BY last_msg DESC LIMIT 50"
    ).fetchall()
    return [dict(r) for r in rows]


# ── Follow-ups ─────────────────────────────────────────────────────────────────

@router.get("/followups")
async def followups(
    filter_type: str = Query("due", description="due | today_sent | failed | paused"),
):
    conn = get_connection()
    now_iso = datetime.now().isoformat()
    today = datetime.now().strftime("%Y-%m-%d")

    if filter_type == "due":
        rows = conn.execute(
            "SELECT fs.*, c.name, c.phone, c.company, c.notes FROM follow_up_schedule fs "
            "JOIN customers c ON fs.customer_id=c.id "
            "WHERE fs.active=1 AND c.status='active' AND fs.next_followup_at <= ? "
            "ORDER BY fs.next_followup_at LIMIT 100", (now_iso,)
        ).fetchall()
    elif filter_type == "today_sent":
        rows = conn.execute(
            "SELECT sf.*, c.name, c.phone FROM sent_followups sf "
            "JOIN customers c ON sf.customer_id=c.id "
            "WHERE sf.status='sent' AND date(sf.sent_at)=? ORDER BY sf.sent_at DESC LIMIT 50",
            (today,)
        ).fetchall()
    elif filter_type == "failed":
        rows = conn.execute(
            "SELECT sf.*, c.name, c.phone FROM sent_followups sf "
            "JOIN customers c ON sf.customer_id=c.id "
            "WHERE sf.status='failed' ORDER BY sf.sent_at DESC LIMIT 50"
        ).fetchall()
    elif filter_type == "paused":
        rows = conn.execute(
            "SELECT fs.*, c.name, c.phone, c.company FROM follow_up_schedule fs "
            "JOIN customers c ON fs.customer_id=c.id "
            "WHERE fs.active=0 ORDER BY fs.id DESC LIMIT 50"
        ).fetchall()
    else:
        rows = []

    return {"filter": filter_type, "items": [dict(r) for r in rows]}


@router.get("/followups/by-date")
async def followups_by_date(date: str = Query(...)):
    conn = get_connection()
    rows = conn.execute(
        "SELECT fs.*, c.name, c.phone, c.company, c.notes FROM follow_up_schedule fs "
        "JOIN customers c ON fs.customer_id=c.id "
        "WHERE fs.active=1 AND c.status='active' AND date(fs.next_followup_at)=? "
        "ORDER BY fs.next_followup_at", (date,)
    ).fetchall()
    return {"date": date, "items": [dict(r) for r in rows]}


@router.post("/followups/{schedule_id}/toggle")
async def toggle_followup(schedule_id: int):
    conn = get_connection()
    row = conn.execute("SELECT active FROM follow_up_schedule WHERE id=?", (schedule_id,)).fetchone()
    if not row:
        return {"error": "not found"}
    new_active = 0 if row["active"] else 1
    conn.execute("UPDATE follow_up_schedule SET active=? WHERE id=?", (new_active, schedule_id))
    conn.commit()
    return {"ok": True, "active": new_active}


@router.post("/followups/{schedule_id}/send-now")
async def send_followup_now(schedule_id: int):
    """Manually trigger a followup message for a schedule."""
    from database import record_followup, save_message, update_followup_schedule
    from followup import generate_followup_message
    from bridge_client import send_message

    conn = get_connection()
    row = conn.execute(
        "SELECT fs.*, c.name, c.phone, c.company, c.notes, c.tags FROM follow_up_schedule fs "
        "JOIN customers c ON fs.customer_id=c.id WHERE fs.id=?", (schedule_id,)
    ).fetchone()
    if not row:
        return {"error": "not found"}

    customer = dict(row)
    message = generate_followup_message(customer)

    result = await send_message(customer["phone"], message)
    cid = customer["customer_id"]

    if result.get("status") == "sent":
        conv_id = save_message(cid, "outbound", message, ai_generated=True)
        record_followup(cid, schedule_id, conv_id, status="sent")
        update_followup_schedule(schedule_id)
        return {"ok": True, "status": "sent", "message": message}
    else:
        record_followup(cid, schedule_id, None, status="failed",
                        error_message=result.get("error", "unknown"))
        return {"ok": False, "status": "failed", "error": result.get("error", "unknown")}


@router.post("/send-manual")
async def send_manual(data: dict):
    phone = data.get("phone", "").strip()
    message = data.get("message", "").strip()
    customer_id = data.get("customer_id")
    if not phone or not message:
        return {"error": "phone and message required"}

    try:
        async with httpx.AsyncClient(timeout=30.0, trust_env=False) as client:
            resp = await client.post("http://127.0.0.1:3001/send", json={"phone": phone, "message": message})
            resp.raise_for_status()
            result = resp.json()
        if customer_id:
            conn = get_connection()
            conn.execute(
                "INSERT INTO conversations(customer_id, direction, content, ai_generated) VALUES(?,?,?,0)",
                (customer_id, "outbound", message)
            )
            conn.commit()
        return {"ok": True, "bridge_result": result}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── Batch Push helpers ────────────────────────────────────────────────────────

@router.post("/followups/add-to-batch")
async def add_customer_to_batch(data: dict):
    """Add a customer to followup schedule for batch push."""
    from database import ensure_followup_schedule
    customer_id = data.get("customer_id")
    if not customer_id:
        return {"error": "customer_id required"}
    conn = get_connection()
    # Ensure customer exists
    cust = conn.execute("SELECT id, name, phone, status FROM customers WHERE id=?", (customer_id,)).fetchone()
    if not cust:
        return {"error": "customer not found"}
    if cust["status"] not in ("active", "inactive"):
        return {"error": f"customer status is {cust['status']}"}
    # If customer is inactive, set back to active
    if cust["status"] == "inactive":
        conn.execute("UPDATE customers SET status='active', updated_at=? WHERE id=?", (datetime.now().isoformat(), customer_id))
        conn.commit()
    # Create/activate followup schedule with next_followup_at=now so it appears immediately
    conn.execute(
        "INSERT INTO follow_up_schedule(customer_id, frequency_days, next_followup_at, active) "
        "VALUES(?, 7, ?, 1) "
        "ON CONFLICT DO NOTHING",
        (customer_id, datetime.now().isoformat())
    )
    # Also ensure any inactive schedule is reactivated
    conn.execute(
        "UPDATE follow_up_schedule SET active=1 WHERE customer_id=? AND active=0",
        (customer_id,)
    )
    conn.commit()
    # Fetch the schedule
    sched = conn.execute(
        "SELECT fs.*, c.name, c.phone, c.company FROM follow_up_schedule fs "
        "JOIN customers c ON fs.customer_id=c.id WHERE fs.customer_id=? AND fs.active=1",
        (customer_id,)
    ).fetchone()
    return {"ok": True, "schedule": dict(sched) if sched else None}


@router.post("/followups/{schedule_id}/remove-from-batch")
async def remove_customer_from_batch(schedule_id: int):
    """Remove a customer from the batch by deactivating their followup schedule."""
    conn = get_connection()
    row = conn.execute("SELECT * FROM follow_up_schedule WHERE id=?", (schedule_id,)).fetchone()
    if not row:
        return {"error": "not found"}
    conn.execute("UPDATE follow_up_schedule SET active=0 WHERE id=?", (schedule_id,))
    conn.commit()
    return {"ok": True}


@router.post("/followups/mark-batch-pushed")
async def mark_batch_pushed():
    """Mark that manual batch push was completed today."""
    from bot_state import mark_batch_pushed_today, _today_str
    mark_batch_pushed_today()
    return {"ok": True, "date": _today_str()}


@router.get("/followups/auto-batch-status")
async def auto_batch_status():
    """Return whether today's batch push was already done manually."""
    from bot_state import was_batch_pushed_today, _today_str
    return {
        "manual_pushed_today": was_batch_pushed_today(),
        "date": _today_str(),
    }

# ── Product Docs ──────────────────────────────────────────────────────────────

@router.get("/product-docs")
async def list_product_docs():
    conn = get_connection()
    rows = conn.execute("SELECT * FROM product_docs ORDER BY updated_at DESC").fetchall()
    return {"items": [dict(r) for r in rows]}


@router.post("/product-docs")
async def create_product_doc(data: dict):
    title = (data.get("title") or "").strip()
    content = (data.get("content") or "").strip()
    if not title:
        return {"error": "title required"}
    conn = get_connection()
    cur = conn.execute(
        "INSERT INTO product_docs(title, content) VALUES(?,?)", (title, content))
    conn.commit()
    row = conn.execute("SELECT * FROM product_docs WHERE id=?", (cur.lastrowid,)).fetchone()
    return {"ok": True, "doc": dict(row)}


@router.put("/product-docs/{doc_id}")
async def update_product_doc(doc_id: int, data: dict):
    conn = get_connection()
    row = conn.execute("SELECT * FROM product_docs WHERE id=?", (doc_id,)).fetchone()
    if not row:
        return {"error": "not found"}
    title = data.get("title", row["title"])
    content = data.get("content", row["content"])
    conn.execute(
        "UPDATE product_docs SET title=?, content=?, updated_at=? WHERE id=?",
        (title, content, datetime.now().isoformat(), doc_id))
    conn.commit()
    row = conn.execute("SELECT * FROM product_docs WHERE id=?", (doc_id,)).fetchone()
    return {"ok": True, "doc": dict(row)}


@router.delete("/product-docs/{doc_id}")
async def delete_product_doc(doc_id: int):
    conn = get_connection()
    conn.execute("DELETE FROM product_docs WHERE id=?", (doc_id,))
    conn.commit()
    return {"ok": True}


@router.post("/product-docs/upload")
async def upload_product_docs(files: list[UploadFile] = File(...)):
    conn = get_connection()
    imported = []
    failed = []

    for file in files:
        filename = file.filename or "untitled"
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        title = filename.rsplit(".", 1)[0] if "." in filename else filename

        try:
            content_bytes = await file.read()
            content = ""

            if ext in ("txt", "md", "markdown"):
                content = content_bytes.decode("utf-8-sig")
            elif ext == "pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content_bytes))
                if reader.is_encrypted:
                    try:
                        reader.decrypt("")
                    except Exception:
                        failed.append({"filename": filename, "reason": "PDF 已加密且无法解密，请提供未加密的版本"})
                        continue
                parts = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
                content = "\n\n".join(parts)
                # OCR fallback for scanned PDFs
                if not content.strip():
                    try:
                        import fitz
                        ocr = _get_ocr()
                        doc = fitz.open(stream=content_bytes, filetype="pdf")
                        ocr_parts = []
                        for i, page in enumerate(doc):
                            pix = page.get_pixmap(dpi=150)
                            img_bytes = pix.tobytes("png")
                            out = ocr(img_bytes)
                            if out.txts:
                                lines = [t.strip() for t in out.txts if t and str(t).strip()]
                                if lines:
                                    ocr_parts.append("\n".join(lines))
                        doc.close()
                        content = "\n\n".join(ocr_parts)
                    except Exception:
                        pass
            elif ext == "docx":
                from docx import Document
                doc = Document(io.BytesIO(content_bytes))
                lines = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        lines.append(p.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            lines.append(row_text)
                content = "\n".join(lines)
                # OCR fallback for scanned/image-based docx
                if not content.strip():
                    try:
                        ocr = _get_ocr()
                        ocr_lines = []
                        for rel in doc.part.rels.values():
                            if "image" in rel.reltype:
                                try:
                                    img_bytes = rel.target_part.blob
                                    out = ocr(img_bytes)
                                    if out.txts:
                                        for t in out.txts:
                                            if t and str(t).strip():
                                                ocr_lines.append(str(t).strip())
                                except Exception:
                                    pass
                        content = "\n".join(ocr_lines)
                    except Exception:
                        pass
            elif ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
                try:
                    ocr = _get_ocr()
                    out = ocr(content_bytes)
                    lines = [str(t).strip() for t in out.txts if t and str(t).strip()]
                    content = "\n".join(lines)
                except Exception as e:
                    failed.append({"filename": filename, "reason": f"图片 OCR 识别失败: {str(e)}"})
                    continue
            elif ext == "xlsx":
                wb = openpyxl.load_workbook(io.BytesIO(content_bytes), read_only=True)
                parts = []
                for name in wb.sheetnames:
                    ws = wb[name]
                    rows = list(ws.iter_rows(values_only=True))
                    if not rows:
                        continue
                    parts.append(f"## {name}")
                    for row in rows:
                        parts.append(" | ".join(str(c or "") for c in row))
                wb.close()
                content = "\n".join(parts)
            elif ext == "xls":
                import xlrd
                wb = xlrd.open_workbook(file_contents=content_bytes)
                parts = []
                for name in wb.sheet_names():
                    ws = wb.sheet_by_name(name)
                    if ws.nrows == 0:
                        continue
                    parts.append(f"## {name}")
                    for r in range(ws.nrows):
                        row_vals = [str(ws.cell_value(r, c) or "") for c in range(ws.ncols)]
                        parts.append(" | ".join(row_vals))
                content = "\n".join(parts)
            elif ext == "csv":
                text = content_bytes.decode("utf-8-sig")
                reader = csv.reader(io.StringIO(text))
                content = "\n".join(" | ".join(row) for row in reader)
            else:
                failed.append({"filename": filename, "reason": f"不支持的文件类型 .{ext}"})
                continue

            if not content.strip():
                failed.append({"filename": filename, "reason": "无法提取文本内容，文件可能为扫描版或空白"})
                continue

            cur = conn.execute(
                "INSERT INTO product_docs(title, content) VALUES(?,?)",
                (title, content[:200000]))
            conn.commit()
            # ── Store original binary for WhatsApp document sending ──
            mime_map = {
                "pdf": "application/pdf",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "xls": "application/vnd.ms-excel",
                "csv": "text/csv",
                "txt": "text/plain", "md": "text/markdown",
                "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp",
            }
            mime_type = mime_map.get(ext, "application/octet-stream")
            from database import save_product_file
            save_product_file(cur.lastrowid, filename, mime_type, content_bytes)
            imported.append({"id": cur.lastrowid, "title": title})
        except Exception as e:
            failed.append({"filename": filename, "reason": f"解析失败: {str(e)}"})

    return {"imported": imported, "failed": failed}


# ── QR Code ────────────────────────────────────────────────────────────────────

@router.get("/qr")
async def get_qr():
    qr = get_pending_qr()
    if qr:
        return {"status": "pending", "qr_text": qr}
    return {"status": "none"}


@router.post("/qr/clear")
async def clear_qr():
    clear_pending_qr()
    return {"ok": True}


# ── Bridge ─────────────────────────────────────────────────────────────────────

@router.get("/bridge")
async def bridge_status():
    status = await _get_bridge_status()
    status["qr_pending"] = get_pending_qr() is not None
    # Read recent bridge log entries
    import os
    log_path = os.path.join(os.path.dirname(__file__), "..", "logs", "bot.log")
    bridge_events = []
    try:
        import subprocess
        result = subprocess.run(
            ["grep", "-E", "WHATSAPP_READY|WHATSAPP_DISCONNECTED|Bridge.*exit|Bridge crashed",
             log_path], capture_output=True, text=True, timeout=5
        )
        for line in result.stdout.strip().split("\n")[-30:]:
            if line.strip():
                bridge_events.append(line.strip())
    except Exception:
        pass

    return {
        **status,
        "recent_events": bridge_events[-20:],
    }


# ── Proxy Config ──────────────────────────────────────────────────────────────

@router.get("/proxy")
async def get_proxy():
    cfg = get_config()
    import yaml
    data_dir = get_data_dir()
    config_path = data_dir / "config.yaml"
    if config_path.exists():
        raw = yaml.safe_load(config_path.read_text(encoding="utf-8"))
        return {"proxy": raw.get("bridge", {}).get("proxy", "")}
    return {"proxy": cfg.bridge.proxy}


@router.post("/proxy")
async def set_proxy(data: dict):
    proxy_url = (data.get("proxy") or "").strip()
    import yaml
    data_dir = get_data_dir()
    config_path = data_dir / "config.yaml"
    if config_path.exists():
        raw = yaml.safe_load(config_path.read_text(encoding="utf-8"))
    else:
        raw = {}
    raw.setdefault("bridge", {})["proxy"] = proxy_url
    config_path.write_text(yaml.dump(raw, allow_unicode=True, default_flow_style=False), encoding="utf-8")
    # Reload config
    from config import load_config
    load_config(str(config_path))
    # Notify bridge to reconnect via env var — bridge needs restart to pick up new proxy
    return {"ok": True, "proxy": proxy_url, "note": "代理已保存。请重启应用或重新扫码以使新代理生效。"}


# ── Bot Pause/Resume ───────────────────────────────────────────────────────────

@router.get("/bot-state")
async def get_bot_state():
    dashboard_conn = get_connection()
    today = datetime.now().strftime("%Y-%m-%d")
    today_replies = dashboard_conn.execute(
        "SELECT COUNT(*) FROM conversations WHERE ai_generated=1 AND direction='outbound' AND date(sent_at)=?",
        (today,)
    ).fetchone()[0]
    today_followups = dashboard_conn.execute(
        "SELECT COUNT(*) FROM sent_followups WHERE status='sent' AND date(sent_at)=?",
        (today,)
    ).fetchone()[0]
    return {
        "paused": is_paused(),
        "today_ai_replies": today_replies,
        "today_followups": today_followups,
    }


@router.post("/bot-state/toggle")
async def toggle_bot_state():
    new_state = toggle_paused()
    return {"paused": new_state, "status": "paused" if new_state else "running"}


async def _get_bridge_status() -> dict:
    cfg = get_config()
    try:
        async with httpx.AsyncClient(timeout=3.0, trust_env=False) as client:
            resp = await client.get(f"http://{cfg.bridge.host}:{cfg.bridge.port}/status")
            bs = resp.json()
            return {
                "online": True,
                "authenticated": bs.get("authenticated", False),
                "phone": bs.get("phone"),
            }
    except Exception as e:
        return {"online": False, "authenticated": False, "phone": None, "error": str(e)}


# ── Agent Monitoring ──────────────────────────────────────────────────────────

@router.get("/agent/status")
async def agent_status():
    """Get agent activity summary: decisions, segment distribution, scores."""
    conn = get_connection()
    today = datetime.now().strftime("%Y-%m-%d")

    # Segment distribution
    segments = conn.execute(
        "SELECT segment, COUNT(*) as cnt FROM lead_scores GROUP BY segment"
    ).fetchall()
    segment_counts = {r["segment"]: r["cnt"] for r in segments}
    total_scored = sum(segment_counts.values())

    # Today's decisions
    today_decisions = conn.execute(
        "SELECT COUNT(*) as cnt FROM agent_decisions WHERE date(created_at)=?",
        (today,)
    ).fetchone()["cnt"]

    # Recent decisions (last 20)
    recent = conn.execute(
        "SELECT ad.*, c.name, c.phone FROM agent_decisions ad "
        "LEFT JOIN customers c ON ad.customer_id = c.id "
        "ORDER BY ad.created_at DESC LIMIT 20"
    ).fetchall()

    # Top hot leads
    hot_leads = conn.execute(
        "SELECT ls.*, c.name, c.phone, c.company FROM lead_scores ls "
        "JOIN customers c ON ls.customer_id = c.id "
        "WHERE ls.segment = 'hot' ORDER BY ls.score DESC LIMIT 10"
    ).fetchall()

    # Memory entries count
    memory_count = conn.execute(
        "SELECT COUNT(*) as cnt FROM ai_memory_entries"
    ).fetchone()["cnt"]

    return {
        "total_scored": total_scored,
        "segment_counts": segment_counts,
        "today_decisions": today_decisions,
        "total_memories": memory_count,
        "hot_leads": [{"id": r["customer_id"], "name": r["name"],
                        "phone": r["phone"], "score": r["score"],
                        "signals": r["signals"]} for r in hot_leads],
        "recent_decisions": [
            {"id": r["id"], "customer_id": r["customer_id"],
             "customer_name": r["name"], "decision_type": r["decision_type"],
             "reasoning": r["reasoning"], "created_at": r["created_at"]}
            for r in recent
        ],
    }


@router.get("/agent/decisions")
async def agent_decisions(limit: int = Query(50, ge=10, le=200)):
    """Get recent agent decision logs."""
    from database import get_recent_decisions
    return {"decisions": get_recent_decisions(limit)}


@router.get("/agent/customer-analysis/{customer_id}")
async def agent_customer_analysis(customer_id: int):
    """Full analysis for a single customer: score, segment, intents, memory."""
    from scoring_engine import score_customer
    from intent_analyzer import analyze_conversation_intent
    from database import get_customer_memory, get_customer_decisions

    score_data = score_customer(customer_id)
    intent_data = analyze_conversation_intent(customer_id)
    memory = get_customer_memory(customer_id, limit=10)
    decisions = get_customer_decisions(customer_id, limit=20)

    # Customer info
    customer = get_connection().execute(
        "SELECT * FROM customers WHERE id=?", (customer_id,)
    ).fetchone()

    return {
        "customer": dict(customer) if customer else None,
        "score": score_data,
        "intent": intent_data,
        "memory": [dict(m) for m in memory],
        "decisions": [dict(d) for d in decisions],
    }


@router.post("/agent/run-scoring")
async def trigger_scoring():
    """Manually trigger a full re-scoring of all customers."""
    from scoring_engine import score_all_customers
    result = score_all_customers()
    return {"ok": True, "result": result}


@router.post("/agent/run-strategy-eval")
async def trigger_strategy_eval():
    """Manually trigger strategy evaluation for all customers."""
    from strategy_manager import evaluate_all_strategies
    result = evaluate_all_strategies()
    return {"ok": True, "result": result}


# ── Activity Log ───────────────────────────────────────────────────────────────

@router.get("/activity")
async def activity_log(lines: int = Query(100, ge=10, le=500)):
    import os, subprocess
    log_path = os.path.join(os.path.dirname(__file__), "..", "logs", "bot.log")
    try:
        result = subprocess.run(
            ["tail", "-n", str(lines), log_path],
            capture_output=True, text=True, timeout=5
        )
        log_lines = result.stdout.strip().split("\n")
        # Filter to INFO and above, skip DEBUG
        filtered = [l for l in log_lines
                    if ("[DEBUG]" not in l
                        and ("[INFO]" in l or "[ERROR]" in l or "[CRITICAL]" in l or "[WARNING]" in l))]
        # Filter out QR code Unicode block-art lines (between QR_RECEIVED and QR_END)
        in_qr = False
        clean = []
        for l in filtered:
            if ">>> QR_RECEIVED <<<" in l:
                in_qr = True
                clean.append(l)
                continue
            if ">>> QR_END <<<" in l:
                in_qr = False
                clean.append(l)
                continue
            if not in_qr:
                clean.append(l)
        return {"lines": clean[-lines:]}
    except Exception as e:
        return {"lines": [], "error": str(e)}


# ── Orders ─────────────────────────────────────────────────────────────────────

@router.get("/orders")
async def orders_list(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=5, le=100),
    search: str = Query(""),
    status: str = Query(""),
    sort: str = Query("created_at"),
    order: str = Query("desc"),
):
    conn = get_connection()
    allowed_sorts = {"order_no", "product", "quantity", "total_amount", "status", "order_date", "created_at"}
    if sort not in allowed_sorts:
        sort = "created_at"
    if order not in ("asc", "desc"):
        order = "desc"

    where = ["1=1"]
    params = []
    if search:
        where.append("(o.order_no LIKE ? OR o.product LIKE ? OR c.name LIKE ? OR c.phone LIKE ?)")
        s = f"%{search}%"
        params.extend([s, s, s, s])
    if status:
        where.append("o.status = ?")
        params.append(status)

    where_clause = " AND ".join(where)

    total = conn.execute(
        f"SELECT COUNT(*) FROM orders o JOIN customers c ON o.customer_id=c.id WHERE {where_clause}",
        params
    ).fetchone()[0]

    params_with_limit = params + [per_page, (page - 1) * per_page]
    rows = conn.execute(
        f"""SELECT o.*, c.name as customer_name, c.phone as customer_phone, c.company as customer_company
            FROM orders o JOIN customers c ON o.customer_id=c.id
            WHERE {where_clause}
            ORDER BY o.{sort} {order} LIMIT ? OFFSET ?""",
        params_with_limit
    ).fetchall()

    return {
        "total": total, "page": page, "per_page": per_page,
        "items": [dict(r) for r in rows],
    }


@router.get("/orders/{order_id}")
async def order_detail(order_id: int):
    conn = get_connection()
    row = conn.execute(
        "SELECT o.*, c.name as customer_name, c.phone as customer_phone, c.company as customer_company "
        "FROM orders o JOIN customers c ON o.customer_id=c.id WHERE o.id=?",
        (order_id,)
    ).fetchone()
    if not row:
        return {"error": "not found"}
    return {"order": dict(row)}


@router.post("/orders")
async def create_order(data: dict):
    conn = get_connection()
    customer_id = data.get("customer_id")
    if not customer_id:
        return {"error": "customer_id required"}
    customer = conn.execute("SELECT id FROM customers WHERE id=?", (customer_id,)).fetchone()
    if not customer:
        return {"error": "customer not found"}

    # Generate order_no: TP + YYYYMMDD + 3-digit seq
    today = datetime.now().strftime("%Y%m%d")
    prefix = f"TP{today}"
    last = conn.execute(
        "SELECT order_no FROM orders WHERE order_no LIKE ? ORDER BY order_no DESC LIMIT 1",
        (f"{prefix}%",)
    ).fetchone()
    if last:
        seq = int(last["order_no"][-3:]) + 1
    else:
        seq = 1
    order_no = f"{prefix}{seq:03d}"

    order_date = data.get("order_date", today)
    delivery_date = data.get("delivery_date", "")
    if order_date and delivery_date and order_date > delivery_date:
        return {"error": "下单日期不能晚于交付日期"}
    now = datetime.now().isoformat()
    conn.execute(
        """INSERT INTO orders(order_no, customer_id, product, quantity, unit_price, total_amount,
           status, order_date, delivery_date, notes, currency, created_at, updated_at)
           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (order_no, customer_id, data.get("product", ""), data.get("quantity", 1),
         data.get("unit_price"), data.get("total_amount"),
         data.get("status", "pending"), order_date,
         delivery_date, data.get("notes", ""), data.get("currency", "USD"), now, now)
    )
    conn.commit()
    row = conn.execute("SELECT o.*, c.name as customer_name FROM orders o JOIN customers c ON o.customer_id=c.id WHERE o.order_no=?", (order_no,)).fetchone()
    return {"ok": True, "order": dict(row)}


@router.put("/orders/{order_id}")
async def update_order(order_id: int, data: dict):
    conn = get_connection()
    row = conn.execute("SELECT * FROM orders WHERE id=?", (order_id,)).fetchone()
    if not row:
        return {"error": "not found"}

    allowed = {"order_no", "customer_id", "product", "quantity", "unit_price", "total_amount",
               "status", "order_date", "delivery_date", "notes", "currency"}
    updates = {k: v for k, v in data.items() if k in allowed}
    if not updates:
        return {"error": "no valid fields"}
    # Validate date order: only fall back to DB if key not in request at all
    new_order_date = updates["order_date"] if "order_date" in updates else row["order_date"]
    new_delivery_date = updates["delivery_date"] if "delivery_date" in updates else row["delivery_date"]
    if new_order_date and new_delivery_date and new_order_date > new_delivery_date:
        return {"error": "下单日期不能晚于交付日期"}
    updates["updated_at"] = datetime.now().isoformat()
    set_clause = ", ".join(f"{k}=?" for k in updates)
    conn.execute(f"UPDATE orders SET {set_clause} WHERE id=?", list(updates.values()) + [order_id])
    conn.commit()
    row = conn.execute("SELECT o.*, c.name as customer_name FROM orders o JOIN customers c ON o.customer_id=c.id WHERE o.id=?", (order_id,)).fetchone()
    return {"ok": True, "order": dict(row)}


@router.delete("/orders/{order_id}")
async def delete_order(order_id: int):
    conn = get_connection()
    row = conn.execute("SELECT id FROM orders WHERE id=?", (order_id,)).fetchone()
    if not row:
        return {"error": "not found"}
    conn.execute("DELETE FROM orders WHERE id=?", (order_id,))
    conn.commit()
    return {"ok": True}


@router.get("/customers/{customer_id}/orders")
async def customer_orders(customer_id: int):
    conn = get_connection()
    rows = conn.execute(
        "SELECT * FROM orders WHERE customer_id=? ORDER BY created_at DESC",
        (customer_id,)
    ).fetchall()
    return [dict(r) for r in rows]
